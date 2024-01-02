from typing import List, Optional, Iterator

from langchain.docstore.document import Document
from langchain.document_loaders.base import BaseLoader
import docx

class DocxLoader(BaseLoader):
    """
    Load a document from an Excel file.
    """
    def __init__(self, path: str, sheet_name: Optional[str] = None, *args, **kwargs):
        self.path = path

    def load(self) -> List[Document]:
        return self.lazy_load()

    def lazy_load(self) -> List[Document]:
        docs = []
        doc = docx.Document(self.path)
        content = []
        for i in range(len(doc.paragraphs)):
            para = doc.paragraphs[i]
            text = para.text
            content.append(text)
        docs.append(
            Document(page_content="".join(content), metadata={"source": self.path})
        )
        return docs
